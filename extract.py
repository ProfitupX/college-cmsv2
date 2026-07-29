import docx
import fitz
import glob
import os

print("Extracting docx files...")
for f in glob.glob('extradata/*.docx'):
    print(f"Reading {f}")
    doc = docx.Document(f)
    text = '\n'.join([p.text for p in doc.paragraphs])
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            text += '\n' + '\t'.join([cell.text for cell in row.cells])
    with open(f + '.txt', 'w', encoding='utf-8') as out:
        out.write(text)

print("Extracting pdf files...")
for f in glob.glob('extradata/*.pdf'):
    print(f"Reading {f}")
    doc = fitz.open(f)
    text = '\n'.join([page.get_text() for page in doc])
    with open(f + '.txt', 'w', encoding='utf-8') as out:
        out.write(text)
print("Done")
